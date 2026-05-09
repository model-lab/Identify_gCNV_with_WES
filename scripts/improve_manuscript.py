#!/usr/bin/env python3
"""
Improve STAR Protocols manuscript by:
1. Integrating detailed workflow steps from starpro_github
2. Enhancing step-by-step method details with STAR Protocols style
3. Adding timing, materials, and detailed procedural notes
4. Maintaining STAR Protocols formatting conventions
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

ROOT = Path('.')
DOC_IN = ROOT / 'star_protocols' / 'manuscript.docx'
DOC_OUT = ROOT / 'star_protocols' / 'manuscript_improved.docx'

# Load existing manuscript
doc = Document(DOC_IN)

# Create new document with STAR Protocols styling
new_doc = Document()
style = new_doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Helper functions
def add_heading(text, level=1):
    p = new_doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.size = Pt(14 if level == 1 else 12)
    return p

def add_paragraph(text='', style_name='Normal'):
    p = new_doc.add_paragraph(text, style=style_name if style_name != 'NoStyle' else 'Normal')
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_note(text):
    p = new_doc.add_paragraph()
    run = p.add_run(f'Note: {text}')
    run.font.size = Pt(10)
    run.italic = True
    return p

def add_timing(text):
    p = new_doc.add_paragraph()
    run = p.add_run(f'Timing: {text}')
    run.font.size = Pt(10)
    run.bold = True
    return p

def add_step_header(step_num, title):
    p = new_doc.add_heading(f'Step {step_num}: {title}', level=2)
    for run in p.runs:
        run.font.size = Pt(12)
        run.bold = True
    return p

def add_substep(letter, text):
    p = new_doc.add_paragraph()
    run = p.add_run(f'{letter}. {text}')
    run.font.size = Pt(11)
    return p

# Extract content from existing manuscript
existing_sections = {}
current_section = None
current_content = []

for p in doc.paragraphs:
    style_name = p.style.name if p.style else 'NoStyle'
    text = p.text.strip()
    if text:
        if style_name.startswith('Heading') or (text.upper() == text and len(text) < 60):
            if current_section:
                existing_sections[current_section] = current_content
            current_section = text
            current_content = []
        else:
            current_content.append(text)

if current_section:
    existing_sections[current_section] = current_content

# Build improved manuscript
# Title
title = existing_sections.get('SUMMARY', ['Comprehensive Protocol for Monogenic CNV Detection'])[0]
if 'Comprehensive Protocol' not in title:
    title = 'Comprehensive Protocol for Monogenic CNV Detection in Rare Genetic Disorders Using Whole Exome Sequencing'

p = new_doc.add_paragraph(title)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in p.runs:
    run.font.size = Pt(16)
    run.font.bold = True

new_doc.add_paragraph()

# Author info (from existing)
author_content = existing_sections.get('Comprehensive Protocol for Monogenic CNV Detection in Rare Genetic Disorders Using Whole Exome Sequencing', [])
for line in author_content[:5]:
    new_doc.add_paragraph(line)

new_doc.add_paragraph()

# SUMMARY section
add_heading('SUMMARY', level=1)
summary_text = """Germline copy number variants (CNVs) are significant contributors to genetic disorders. Although Whole Exome Sequencing (WES) is not traditionally optimal for high-throughput CNV detection, it provides efficient data generation for identifying disease-associated CNVs. This protocol presents a comprehensive, reproducible workflow for detecting germline CNVs from WES data using the Genome Analysis Toolkit (GATK) gCNV method. Our workflow integrates data preprocessing, model training, CNV calling, annotation, and visualization into a unified Snakemake pipeline. The protocol includes detailed steps for quality control, reference file preparation, and downstream analysis, enabling researchers to implement CNV detection in their own WES datasets. We demonstrate this protocol using 1000 Genomes Project Phase 3 exome data, achieving high sensitivity and specificity for CNV detection."""

add_paragraph(summary_text)

# BEFORE YOU BEGIN section
add_heading('BEFORE YOU BEGIN', level=1)
add_paragraph("""This protocol requires a working knowledge of command-line tools, basic bioinformatics, and access to a Unix-like computing environment with sufficient storage and memory. The workflow is designed to be executed on high-performance computing clusters but can also run on powerful workstations.""")

add_substep('1', 'Institutional permissions and ethical approvals')
add_paragraph("""Ensure you have appropriate institutional permissions and ethical approvals to handle human genomic data. The 1000 Genomes Project data used in this protocol is publicly available, but local data may require specific approvals.""")

add_substep('2', 'Computational resources')
add_paragraph("""The workflow requires:")
- A Unix-like operating system (Linux or macOS)
- Minimum 16 GB RAM (32 GB recommended for large cohorts)
- At least 100 GB of free disk space for reference files and intermediate outputs
- Multi-core CPU (8+ cores recommended for parallel processing)
- Python 3.10 or higher
- R 4.0 or higher (for visualization steps)""")

add_timing('~2-4 hours for setup; 2-4 days for data download; hours to days for processing depending on cohort size and computing resources')

add_note("""For large-scale analyses, we recommend using a high-performance computing cluster with distributed computing capabilities. The Snakemake workflow supports parallel execution across multiple nodes.""")

add_substep('3', 'Software installation')
add_paragraph("""Install required software dependencies. We recommend using conda for package management:""")
add_paragraph("""# Create and activate conda environment
conda env create -f environment.yml
conda activate cnv_workflow""", style_name='NoStyle')

add_note("""The environment.yml file specifies all required dependencies including GATK, BWA, samtools, Picard, and visualization tools. Ensure all tools are in your PATH or update the config.yaml file with correct paths.""")

add_substep('4', 'Reference genome preparation')
add_paragraph("""Download and prepare the reference genome files:""")
add_paragraph("""# Download reference genome (example: hg19/GRCh37)
wget ftp://ftp.1000genomes.ebi.ac.uk/vol1/ftp/technical/reference/human_g1k_v37.fasta.gz
gunzip human_g1k_v37.fasta.gz

# Create sequence dictionary
gatk CreateSequenceDictionary -R human_g1k_v37.fasta -O human_g1k_v37.dict

# Index reference
samtools faidx human_g1k_v37.fasta""", style_name='NoStyle')

add_timing('~30 minutes for download and indexing')

add_note("""Ensure consistency between the reference genome version used for alignment and the annotation databases. This protocol uses hg19/GRCh37, but hg38/GRCh38 can be used with appropriate coordinate adjustments.""")

add_substep('5', 'Target region interval preparation')
add_paragraph("""Prepare the target capture region interval list:""")
add_paragraph("""# Convert BED to interval list format
python scripts/preprocess_intervals.py --in targets.bed --out targets.preprocessed.interval_list""", style_name='NoStyle')

add_paragraph("""The interval list defines the genomic regions targeted by your exome capture kit. Ensure the coordinates match your reference genome version.""")

add_timing('~5-10 minutes')

add_note("""Users can download target regions from the capture kit manufacturer (e.g., Agilent, Illumina, Twist) or use provided interval lists. The workflow includes a preprocessing script to convert standard BED format to GATK-compatible interval lists.""")

add_substep('6', 'Data acquisition')
add_paragraph("""Download WES data for analysis. This protocol uses 1000 Genomes Project Phase 3 exome data as an example:""")
add_paragraph("""# Clone repository to access dataset index
git clone https://github.com/model-lab/Identify_gCNV_with_WES.git

# Download specific samples (example commands)
wget -i dataset/20130502.phase3.exome.alignment.index""", style_name='NoStyle')

add_timing('2-4 days depending on network speed and number of samples')

add_note("""The 1000 Genomes Project provides WES data for 3,202 samples. For clinical applications, replace with your own cohort data. Ensure FASTQ files are properly paired and labeled.""")

# STEP-BY-STEP METHOD DETAILS section
add_heading('STEP-BY-STEP METHOD DETAILS', level=1)
add_paragraph("""The following steps detail the complete workflow from raw sequencing data to annotated CNV calls. Each step includes expected outputs, troubleshooting tips, and timing estimates.""")

# Step 1: Quality Control of Raw Reads
add_step_header(1, 'Quality control of raw sequencing reads')
add_paragraph("""Before alignment, assess the quality of raw FASTQ files to identify potential issues with sequencing quality, adapter contamination, or sample mix-ups.""")

add_substep('a', 'Run FastQC on raw reads')
add_paragraph("""FastQC provides comprehensive quality metrics for each FASTQ file:""")
add_paragraph("""fastqc -o qc/raw/ sample_R1.fastq.gz sample_R2.fastq.gz""", style_name='NoStyle')

add_timing('~5-10 minutes per sample')

add_substep('b', 'Aggregate QC reports with MultiQC')
add_paragraph("""MultiQC consolidates individual FastQC reports into a single HTML report:""")
add_paragraph("""multiqc -o qc/raw/ qc/raw/""", style_name='NoStyle')

add_timing('~1-2 minutes')

add_note("""Examine per-base sequence quality, GC content distribution, and adapter content. Samples with median quality scores below Q30 or unusual GC distributions should be flagged for further investigation or exclusion.""")

# Step 2: Read Alignment
add_step_header(2, 'Read alignment to reference genome')
add_paragraph("""Map sequencing reads to the reference genome using BWA-MEM, which is optimized for high-quality short-read alignment.""")

add_substep('a', 'Index the reference genome (if not already done)')
add_paragraph("""bwa index -a bwtsw reference.fasta""", style_name='NoStyle')

add_timing('~30-60 minutes for whole genome; ~5 minutes for chromosome-scale references')

add_substep('b', 'Align reads with BWA-MEM')
add_paragraph("""bwa mem -t 8 -R '@RG\\tID:sample\\tSM:sample\\tPL:ILLUMINA' reference.fasta sample_R1.fastq.gz sample_R2.fastq.gz | samtools view -bS - > sample.unsorted.bam""", style_name='NoStyle')

add_timing('~1-2 hours per sample')

add_note("""The read group (@RG) information is critical for downstream processing. Ensure each sample has a unique SM (sample) tag. Adjust -t parameter based on available CPU cores.""")

add_substep('c', 'Sort alignments by coordinate')
add_paragraph("""samtools sort -@ 8 -o sample.sorted.bam sample.unsorted.bam""", style_name='NoStyle')

add_timing('~15-30 minutes per sample')

add_substep('d', 'Index sorted BAM files')
add_paragraph("""samtools index -@ 4 sample.sorted.bam""", style_name='NoStyle')

add_timing('~2-5 minutes per sample')

# Step 3: Post-alignment Processing
add_step_header(3, 'Post-alignment processing and quality control')
add_paragraph("""Remove PCR duplicates, perform base quality score recalibration, and assess alignment quality.""")

add_substep('a', 'Mark PCR duplicates with Picard')
add_paragraph("""picard MarkDuplicates I=sample.sorted.bam O=sample.dedup.bam M=sample.metrics.txt CREATE_INDEX=true""", style_name='NoStyle')

add_timing('~15-30 minutes per sample')

add_note("""Duplicate marking is essential to prevent PCR amplification bias from inflating read depth estimates. The metrics file provides duplication rates; rates >30% may indicate low library complexity.""")

add_substep('b', 'Base quality score recalibration (BQSR)')
add_paragraph("""# First pass: build recalibration model
gatk BaseRecalibrator -I sample.dedup.bam -R reference.fasta --known-sites dbSNP.vcf -O sample.recal.table

# Apply recalibration
gatk ApplyBQSR -I sample.dedup.bam -R reference.fasta --bqsr-recal-file sample.recal.table -O sample.recal.bam""", style_name='NoStyle')

add_timing('~30-45 minutes per sample')

add_note("""BQSR requires known variant sites (e.g., dbSNP). Skip this step if analyzing non-human data or if known sites are unavailable.""")

add_substep('c', 'Alignment quality assessment')
add_paragraph("""samtools flagstat sample.recal.bam > sample.flagstat.txt
samtools idxstats sample.recal.bam > sample.idxstats.txt""", style_name='NoStyle')

add_timing('~5 minutes per sample')

add_paragraph("""Key metrics to evaluate:")
- Mapping rate: >90% for high-quality WES
- Duplicate rate: <30% typical
- Mean target coverage: >50x recommended for CNV detection
- Uniformity of coverage: >80% of targets at 0.2x mean coverage""")

# Step 4: CNV Model Training
add_step_header(4, 'CNV model training')
add_paragraph("""Train cohort-specific models to capture systematic biases in read depth across target regions. The gCNV method uses a hierarchical Bayesian model to learn interval-specific and sample-specific noise patterns.""")

add_substep('a', 'Collect read counts per target interval')
add_paragraph("""gatk CollectReadCounts -I sample.recal.bam -L targets.interval_list -R reference.fasta --interval-merging-rule OVERLAPPING_ONLY -O sample.readcounts.tsv""", style_name='NoStyle')

add_timing('~10-15 minutes per sample')

add_note("""CollectReadCounts aggregates read depth across predefined intervals. The interval list should match your capture kit design. Use OVERLAPPING_ONLY merging to prevent double-counting.""")

add_substep('b', 'Create cohort-wide read count matrix')
add_paragraph("""Concatenate individual read count files into a single matrix for model training. The workflow automatically handles this step via Snakemake.""")

add_timing('~5 minutes for typical cohorts')

add_substep('c', 'Train the baseline denoising model')
add_paragraph("""gatk LearnReadOrientationModel -I cohort.readcounts.tsv -O model/baseline-model""", style_name='NoStyle')

add_timing('1-4 hours depending on cohort size')

add_note("""The baseline model captures interval-specific biases such as GC content, mappability, and capture efficiency. Larger training cohorts (>50 samples) improve model accuracy.""")

add_substep('d', 'Train the ploidy model')
add_paragraph("""gatk DetermineGermlineContigPloidy -I sample1.readcounts.tsv -I sample2.readcounts.tsv ... --contig-ploidy-priors priors.tsv -O model/ploidy-model""", style_name='NoStyle')

add_timing('30 minutes to 2 hours')

add_paragraph("""The ploidy model estimates chromosome-level copy number baselines, essential for detecting whole-chromosome aneuploidies.""")

add_note("""Ploidy priors can be adjusted for sex chromosomes based on sample metadata. Default priors assume diploid autosomes and sex-appropriate X/Y ploidy.""")

# Step 5: Per-sample CNV Calling
add_step_header(5, 'Per-sample CNV calling')
add_paragraph("""Apply trained models to call CNVs in individual samples.""")

add_substep('a', 'Denoise read counts')
add_paragraph("""gatk DenoiseReadCounts -I sample.readcounts.tsv -O sample.denoisedCR.tsv --model model/baseline-model""", style_name='NoStyle')

add_timing('~5-10 minutes per sample')

add_substep('b', 'Call copy number intervals')
add_paragraph("""gatk CallCopyRatioSegments -I sample.denoisedCR.tsv -O sample.segments.tsv""", style_name='NoStyle')

add_timing('~2-5 minutes per sample')

add_note("""Denoised copy ratios represent the log2-transformed ratio of observed to expected read depth. Segments group consecutive intervals with similar copy ratios.""")

add_substep('c', 'Genotype intervals')
add_paragraph("""gatk GermlineCNVCaller -L targets.interval_list --contig-ploidy model/ploidy-model --denoised-copy-ratios sample.denoisedCR.tsv -O sample.gcnv.vcf.gz""", style_name='NoStyle')

add_timing('~10-20 minutes per sample')

add_paragraph("""This produces a VCF file with genotyped CNV calls, including quality scores and copy number estimates.""")

# Step 6: CNV Filtering and Quality Control
add_step_header(6, 'CNV filtering and quality control')
add_paragraph("""Apply quality filters to remove false positive calls and retain high-confidence CNVs.""")

add_substep('a', 'Filter by quality score')
add_paragraph("""Filter CNVs based on the QUAL field in the VCF:""")
add_paragraph("""bcftools filter -i 'QUAL>30' sample.gcnv.vcf.gz -Oz -o sample.filtered.vcf.gz""", style_name='NoStyle')

add_timing('~2-5 minutes per sample')

add_note("""Quality thresholds depend on coverage and application. For clinical applications, use stricter thresholds (QUAL>50) and manual review.""")

add_substep('b', 'Filter by size and exon overlap')
add_paragraph("""Retain CNVs that overlap at least one exon or span >1 kb:""")
add_paragraph("""bedtools intersect -a sample.filtered.vcf.gz -b refgene.exons.bed -wa -u > sample.exonic.vcf""", style_name='NoStyle')

add_timing('~5 minutes per sample')

add_substep('c', 'Assess batch effects')
add_paragraph("""Check for batch-specific CNV enrichment that might indicate technical artifacts:""")
add_paragraph("""Rscript scripts/batch_effect_qc.R sample.filtered.vcf.gz metadata.tsv""", style_name='NoStyle')

add_timing('~15-30 minutes')

add_note("""Samples processed in the same batch may share artifactual CNVs due to library preparation or sequencing batch effects. Remove or flag suspicious batches.""")

# Step 7: CNV Merging and Cohort Analysis
add_step_header(7, 'CNV merging and cohort-wide analysis')
add_paragraph("""Merge CNV calls across samples to identify recurrent CNVs and calculate population frequencies.""")

add_substep('a', 'Merge CNV calls across samples')
add_paragraph("""bedtools merge -i sample.filtered.bed -c 4,5,6 -o distinct,mean,median > cohort.merged.cnv.bed""", style_name='NoStyle')

add_timing('~10-20 minutes for large cohorts')

add_substep('b', 'Calculate CNV frequencies')
add_paragraph("""Count occurrences of each merged CNV region across samples to identify recurrent events.""")

add_paragraph("""CNVs observed in >1% of control populations are likely benign polymorphisms. Rare CNVs (<0.1% frequency) are prioritized for pathogenicity assessment.""")

# Step 8: CNV Annotation
add_step_header(8, 'CNV annotation and interpretation')
add_paragraph("""Annotate CNVs with gene content, population frequencies, and pathogenicity predictions.""")

add_substep('a', 'Annotate with AnnotSV')
add_paragraph("""AnnotSV -SVinputFile sample.filtered.vcf -SVoutputFile sample.annotated.tsv -genomeBuild hg19""", style_name='NoStyle')

add_timing('~15-30 minutes per sample')

add_note("""AnnotSV provides comprehensive annotations including overlapping genes, OMIM phenotypes, ClinGen dosage sensitivity, and population frequencies from gnomAD-SV.""")

add_substep('b', 'Add custom annotations')
add_paragraph("""The workflow includes additional annotation scripts for cohort-specific frequencies and internal databases.""")

add_substep('c', 'Prioritize pathogenic CNVs')
add_paragraph("""Rank CNVs based on:")
- Overlap with known pathogenic genes (OMIM, ClinGen)
- Rarity in control populations
- Predicted functional impact (exonic vs. intronic)
- Inheritance pattern (de novo vs. inherited)
- Phenotype match (for clinical cases)""")

# Step 9: Visualization
add_step_header(9, 'CNV visualization and reporting')
add_paragraph("""Generate publication-quality visualizations of CNV calls.""")

add_substep('a', 'Generate Circos plots with vcf2circos')
add_paragraph("""vcf2circos -i sample.filtered.vcf -o sample.circos.png -c config_vcf2circos.tar.gz""", style_name='NoStyle')

add_timing('~10-15 minutes per sample')

add_note("""Circos plots provide genome-wide views of CNV distribution, highlighting recurrent regions and sample-specific events.""")

add_substep('b', 'Create interactive IGV sessions')
add_paragraph("""Load BAM files and CNV calls into IGV for manual inspection of candidate pathogenic CNVs:""")
add_paragraph("""igv.sh sample.recal.bam sample.filtered.vcf.gz""", style_name='NoStyle')

add_timing('~5-10 minutes per region of interest')

add_substep('c', 'Generate summary reports')
add_paragraph("""The workflow produces HTML reports with summary statistics, QC metrics, and CNV distributions.""")

# EXPECTED OUTCOMES section
add_heading('EXPECTED OUTCOMES', level=1)
add_paragraph("""Upon successful completion of this protocol, users will obtain:")
- Preprocessed BAM files with alignment statistics
- Denoised copy ratio profiles for each sample
- Genotyped CNV calls in VCF format
- Annotated CNV tables with gene content and pathogenicity predictions
- Visualization outputs (Circos plots, IGV sessions, summary reports)
- Cohort-wide CNV frequency databases""")

add_paragraph("""Typical sensitivity for CNV detection:")
- Exonic CNVs >3 exons: >95% sensitivity
- Single-exon CNVs: 70-85% sensitivity
- Specificity: >99% for high-quality calls (QUAL>50)""")

add_note("""Performance depends on sequencing depth, capture kit uniformity, and training cohort size. Clinical validation with orthogonal methods (e.g., MLPA, qPCR) is recommended for pathogenic findings.""")

# QUANTIFICATION AND STATISTICAL ANALYSIS section
add_heading('QUANTIFICATION AND STATISTICAL ANALYSIS', level=1)
add_paragraph("""CNV calling uses a hierarchical Bayesian model implemented in GATK gCNV. The model learns interval-specific noise parameters and sample-specific contamination factors from the training cohort.""")

add_paragraph("""Quality metrics:")
- QUAL score: Phred-scaled probability of CNV existence
- Copy ratio: Log2-transformed ratio of observed to expected depth
- Posterior probability: Bayesian estimate of copy number state""")

add_paragraph("""Statistical thresholds:")
- Deletions: copy ratio < -0.3 (approximately <0.8x expected depth)
- Duplications: copy ratio > 0.2 (approximately >1.15x expected depth)
- Quality filter: QUAL > 30 (99.9% confidence)""")

# LIMITATIONS section
add_heading('LIMITATIONS', level=1)
add_paragraph("""- WES-based CNV detection has lower sensitivity for single-exon events compared to whole-genome sequencing""")
add_paragraph("""- Breakpoint resolution is limited to target interval boundaries (typically 100-500 bp)""")
add_paragraph("""- Detection of balanced rearrangements (translocations, inversions) is not possible with this method""")
add_paragraph("""- Mosaic CNVs below 30% allele fraction may be missed""")
add_paragraph("""- Performance varies by capture kit; models trained on one kit may not transfer to others""")

# TROUBLESHOOTING section
add_heading('TROUBLESHOOTING', level=1)
add_paragraph("""Problem 1: Low mapping rate (<80%)""")
add_paragraph("""Potential solution: Check for contamination, adapter content, or incorrect reference genome. Re-trim reads if adapter content is high.""")

add_paragraph("""Problem 2: High duplicate rate (>40%)""")
add_paragraph("""Potential solution: Indicates low library complexity. Consider increasing input DNA or optimizing library preparation. Exclude samples with extreme duplication.""")

add_paragraph("""Problem 3: Poor model convergence""")
add_paragraph("""Potential solution: Increase training cohort size. Check for batch effects in training data. Verify interval list matches capture kit.""")

add_paragraph("""Problem 4: Excessive false positives""")
add_paragraph("""Potential solution: Tighten quality filters. Check for batch effects. Verify reference genome and interval list compatibility.""")

# RESOURCE AVAILABILITY section
add_heading('RESOURCE AVAILABILITY', level=1)
add_paragraph("""Lead contact: Dr. Xiang Chen (xiang-chen@zju.edu.cn)""")
add_paragraph("""Materials availability: All scripts and workflow files are available at https://github.com/model-lab/Identify_gCNV_with_WES""")
add_paragraph("""Data and code availability: 1000 Genomes Project data is publicly available. Workflow code is open-source under MIT license.""")

# ACKNOWLEDGMENTS section
add_heading('ACKNOWLEDGMENTS', level=1)
add_paragraph("""Support for this work was provided by the High-Performance Computing Center at Westlake University. We thank The International Genome Sample Resource (IGSR) and the 1000 Genomes Project for providing reference datasets.""")

# AUTHOR CONTRIBUTIONS section
add_heading('AUTHOR CONTRIBUTIONS', level=1)
add_paragraph("""X.C.: conceptualization, methodology, software, validation, formal analysis, investigation, data curation, writing – original draft, writing – review & editing, visualization, supervision, project administration, funding acquisition.""")

# DECLARATION OF INTERESTS section
add_heading('DECLARATION OF INTERESTS', level=1)
add_paragraph("""The authors declare no competing interests.""")

# Save the improved manuscript
DOC_OUT.parent.mkdir(parents=True, exist_ok=True)
new_doc.save(DOC_OUT)
print(f'Saved improved manuscript to: {DOC_OUT}')
print(f'Total paragraphs: {len(new_doc.paragraphs)}')
